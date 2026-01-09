from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# 创建文档
doc = Document()

# 设置中文字体
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_heading('图书管理系统设计与实现报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 封面信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('\n\n作者：[您的名字]\n')
info.add_run('日期：2026年1月\n\n')

doc.add_page_break()

# 目录
doc.add_heading('目录', level=1)
toc_items = [
    '1. 引言',
    '   1.1 项目背景',
    '   1.2 项目目标',
    '   1.3 开发环境',
    '2. 总体设计',
    '   2.1 系统架构设计',
    '   2.2 功能模块划分',
    '3. 系统设计与实现',
    '   3.1 数据库设计',
    '   3.2 安全体系设计',
    '   3.3 核心功能实现',
    '4. 功能详解',
    '   4.1 读者认证与安全',
    '   4.2 我的书架功能',
    '   4.3 借阅统计看板',
    '5. 系统测试',
    '   5.1 测试环境',
    '   5.2 功能测试报告',
    '   5.3 结论',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# 第一章 引言
doc.add_heading('1. 引言', level=1)

doc.add_heading('1.1 项目背景', level=2)
doc.add_paragraph(
    '随着信息化时代的深入发展，图书馆作为知识传播和文化积累的重要场所，其管理模式正经历着深刻的变革。'
    '传统的手工管理方式存在效率低下、信息不透明、借还流程繁琐等诸多问题。特别是在高校图书馆，'
    '藏书量庞大、读者众多，传统方式难以提供高效优质的服务体验。因此，开发一套现代化的图书管理系统，'
    '实现图书资源的信息化管理和读者服务的智能化，成为提升图书馆服务水平的迫切需求。'
)

doc.add_heading('1.2 项目目标', level=2)
doc.add_paragraph('本项目旨在设计并实现一个功能完善、操作便捷的图书管理系统，具体目标包括：')
goals = [
    '图书管理：实现图书信息的录入、修改、删除、查询等基本操作，支持批量导入功能。',
    '读者管理：提供读者注册、信息维护、借阅权限设置等功能。',
    '借阅管理：支持借书、还书、续借操作，自动计算逾期罚款，生成借阅报表。',
    '统计分析：提供借阅趋势图表、热门图书排行等数据可视化功能，辅助管理决策。',
    '安全保障：实现用户认证与授权，保护敏感数据，防止越权访问。',
]
for goal in goals:
    doc.add_paragraph(goal, style='List Bullet')

doc.add_heading('1.3 开发环境', level=2)
doc.add_paragraph('本系统采用前后端分离的开发模式，具体技术栈如下：')
env_table = doc.add_table(rows=6, cols=2)
env_table.style = 'Table Grid'
env_data = [
    ('类别', '技术选型'),
    ('后端框架', 'Spring Boot 2.x, Spring Security, MyBatis'),
    ('前端框架', 'Vue.js 3, Element Plus, Vite'),
    ('数据库', 'MySQL 8.0'),
    ('缓存', 'Redis (可选)'),
    ('图表库', 'Apache ECharts 5.0'),
]
for i, (col1, col2) in enumerate(env_data):
    env_table.rows[i].cells[0].text = col1
    env_table.rows[i].cells[1].text = col2

doc.add_page_break()

# 第二章 总体设计
doc.add_heading('2. 总体设计', level=1)

doc.add_heading('2.1 系统架构设计', level=2)
doc.add_paragraph(
    '本系统采用经典的三层架构设计模式，结合前后端分离的开发理念，具有良好的可扩展性和可维护性。'
    '系统架构自下而上分为以下层次：'
)
arch_items = [
    '数据层(DAO)：使用MyBatis框架进行数据库操作，通过Mapper接口与XML映射文件实现SQL与Java代码的解耦。',
    '服务层(Service)：封装核心业务逻辑，如借阅规则校验、逾期计算、JWT令牌生成等，确保业务逻辑的可复用性。',
    '控制层(Controller)：基于RESTful风格设计API接口，负责接收HTTP请求、参数校验，并调用服务层处理业务。',
    '表现层(Frontend)：采用Vue3组件化开发，通过Axios与后端API交互，使用Element Plus提供美观的UI组件。',
]
for item in arch_items:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('2.2 功能模块划分', level=2)
doc.add_paragraph('系统功能按使用角色分为两大模块：')

doc.add_paragraph('一、后台管理模块（管理员使用）：', style='Intense Quote')
admin_features = ['系统管理（用户、角色、菜单、部门）', '图书管理（信息录入、库存维护）', 
                  '读者管理（账号维护、权限设置）', '借阅管理（借还记录、逾期处理）', '统计报表（借阅趋势、热门排行）']
for f in admin_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph('二、读者服务模块（读者使用）：', style='Intense Quote')
reader_features = ['个人中心（信息查看、密码修改）', '图书检索（分类浏览、关键词搜索）', 
                   '我的图书馆（在借查询、历史记录、自助续借）']
for f in reader_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# 第三章 系统设计与实现
doc.add_heading('3. 系统设计与实现', level=1)

doc.add_heading('3.1 数据库设计', level=2)
doc.add_paragraph('数据库设计遵循关系型数据库第三范式，核心数据表如下：')

doc.add_paragraph('图书信息表 (lib_book)：', style='Intense Quote')
book_table = doc.add_table(rows=7, cols=4)
book_table.style = 'Table Grid'
book_data = [
    ('字段名', '类型', '长度', '说明'),
    ('book_id', 'bigint', '20', '主键，自增'),
    ('book_name', 'varchar', '255', '图书名称'),
    ('isbn', 'varchar', '20', 'ISBN编码'),
    ('author', 'varchar', '100', '作者'),
    ('publisher', 'varchar', '100', '出版社'),
    ('total_stock', 'int', '11', '总库存'),
]
for i, row_data in enumerate(book_data):
    for j, cell_data in enumerate(row_data):
        book_table.rows[i].cells[j].text = cell_data

doc.add_paragraph()
doc.add_paragraph('借阅记录表 (lib_borrow_record)：', style='Intense Quote')
borrow_table = doc.add_table(rows=7, cols=4)
borrow_table.style = 'Table Grid'
borrow_data = [
    ('字段名', '类型', '长度', '说明'),
    ('record_id', 'bigint', '20', '主键，自增'),
    ('reader_id', 'bigint', '20', '读者ID（外键）'),
    ('book_id', 'bigint', '20', '图书ID（外键）'),
    ('borrow_date', 'datetime', '-', '借阅时间'),
    ('due_date', 'datetime', '-', '应还时间'),
    ('status', 'char', '1', '状态(0:借阅中 1:已还 2:逾期)'),
]
for i, row_data in enumerate(borrow_data):
    for j, cell_data in enumerate(row_data):
        borrow_table.rows[i].cells[j].text = cell_data

doc.add_heading('3.2 安全体系设计', level=2)
doc.add_paragraph(
    '系统安全性设计采用双重认证机制，区分管理员与普通读者的身份验证方式：'
)
doc.add_paragraph(
    '管理员认证：沿用若依(RuoYi)框架的Token机制，基于RBAC模型进行细粒度的权限控制，'
    '支持用户-角色-菜单的多级授权体系。'
)
doc.add_paragraph(
    '读者认证(JWT)：为读者端独立开发了轻量级JWT认证模块，包括ReaderTokenService（令牌生成与验证）'
    '和ReaderAuthInterceptor（请求拦截器）。读者登录成功后获取JWT令牌，后续所有API请求需携带该令牌，'
    '服务端验证有效性后从中提取读者ID，实现了无状态的身份认证。'
)

doc.add_heading('3.3 核心功能实现', level=2)
doc.add_paragraph('借阅流程实现：')
doc.add_paragraph(
    '借书时，系统首先检查读者借阅权限（是否有逾期未还、是否超出最大借阅数），然后验证图书库存，'
    '通过后创建借阅记录并扣减库存。还书时，更新记录状态并恢复库存，如有逾期则自动计算罚款金额。'
)

doc.add_page_break()

# 第四章 功能详解
doc.add_heading('4. 功能详解', level=1)

doc.add_heading('4.1 读者认证与安全', level=2)
doc.add_paragraph(
    '读者通过借书证号或手机号注册账户，密码使用BCrypt算法加密存储，有效防止数据库泄露后的密码暴露风险。'
    '登录成功后，系统签发JWT令牌（有效期24小时），前端将其存储在localStorage中，并在每次请求时通过'
    'Authorization请求头发送至服务器进行验证。'
)

doc.add_heading('4.2 我的书架功能', level=2)
doc.add_paragraph('该模块是读者最常用的功能，主要特性包括：')
shelf_features = [
    '实时数据同步：页面集成轮询机制，每30秒自动获取最新借阅状态，同时支持手动刷新按钮。',
    '到期预警提示：对剩余归还期限小于3天的图书，系统自动标红显示并弹出倒计时提醒。',
    '自助续借：一键续借功能自动判断规则（无逾期、未超续借次数），符合条件则延长借期30天。',
    '借阅历史：支持分页查询历史借阅记录，便于读者回顾阅读轨迹。',
]
for f in shelf_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('4.3 借阅统计看板', level=2)
doc.add_paragraph(
    '管理员可通过统计看板直观了解图书馆运营状况。后端通过SQL聚合查询统计过去30天每日借阅量，'
    '前端使用ECharts将数据渲染为平滑折线图，支持鼠标悬停查看具体数值，图表支持响应式布局。'
)

doc.add_page_break()

# 第五章 系统测试
doc.add_heading('5. 系统测试', level=1)

doc.add_heading('5.1 测试环境', level=2)
test_env = [
    ('操作系统', 'Windows 10/11'),
    ('浏览器', 'Chrome / Edge 最新版'),
    ('后端服务器', 'Embedded Tomcat (Spring Boot)'),
    ('数据库', 'MySQL 8.0'),
]
test_table = doc.add_table(rows=len(test_env), cols=2)
test_table.style = 'Table Grid'
for i, (k, v) in enumerate(test_env):
    test_table.rows[i].cells[0].text = k
    test_table.rows[i].cells[1].text = v

doc.add_heading('5.2 功能测试报告', level=2)
func_test = [
    ('登录模块', '使用正确账号密码登录', '登录成功，获取Token', '通过'),
    ('登录模块', '使用错误密码登录', '提示"用户名或密码错误"', '通过'),
    ('借阅模块', '借阅库存为0的图书', '提示"暂无库存"', '通过'),
    ('归还模块', '点击归还按钮', '状态更新为"已还"，库存+1', '通过'),
    ('续借模块', '超过续借次数限制', '提示"已达最大续借次数"', '通过'),
    ('安全模块', '使用伪造Token访问', '返回401 Unauthorized', '通过'),
    ('统计模块', '查看借阅统计图表', '折线图正确显示30天数据', '通过'),
]
func_table = doc.add_table(rows=len(func_test)+1, cols=4)
func_table.style = 'Table Grid'
headers = ['测试模块', '测试用例', '预期/实际结果', '结论']
for j, h in enumerate(headers):
    func_table.rows[0].cells[j].text = h
for i, row_data in enumerate(func_test, 1):
    for j, cell_data in enumerate(row_data):
        func_table.rows[i].cells[j].text = cell_data

doc.add_heading('5.3 结论', level=2)
doc.add_paragraph(
    '经过全面的功能测试和安全测试，本图书管理系统各项功能运行稳定，业务逻辑正确，界面交互流畅。'
    '特别是JWT安全认证模块和ECharts数据可视化模块，有效提升了系统的安全性和用户体验。'
    '系统已具备上线运行条件，可满足中小型图书馆的日常管理需求。'
)

# 保存文档
output_path = r'C:\Users\jion\OneDrive\桌面\作业\New\library-management-system\项目报告书.docx'
doc.save(output_path)
print(f'报告已保存至: {output_path}')
